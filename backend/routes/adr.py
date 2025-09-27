from flask import Blueprint, request, jsonify, session, send_file
from backend.database import get_db
from datetime import datetime
import io
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors

adr_bp = Blueprint("adr_api", __name__)


@adr_bp.route("/api/adr-report", methods=["POST"])
def submit_adr_report():
    data = request.get_json()
    if not data:
        return jsonify({"error": "Invalid data"}), 400

    required_fields = ["drug_id", "reaction_description"]
    if not all(field in data for field in required_fields):
        return jsonify({"error": "Missing required fields"}), 400

    conn = get_db()
    conn.execute(
        """
        INSERT INTO adr_reports (
            drug_id, patient_age_range, patient_gender,
            reaction_description, reaction_start_date, other_medications, status
        ) VALUES (?, ?, ?, ?, ?, ?, ?)
        """,
        (
            data["drug_id"],
            data.get("patient_age_range"),
            data.get("patient_gender"),
            data["reaction_description"],
            data.get("reaction_start_date"),
            data.get("other_medications"),
            "New",
        ),
    )
    conn.commit()

    return jsonify({"message": "Adverse Drug Reaction report submitted successfully."}), 201


@adr_bp.route("/api/adr-report/<int:report_id>", methods=["DELETE"])
def delete_adr_report(report_id):
    if not session.get("admin_id"):
        return jsonify({"error": "Authentication required"}), 401

    try:
        conn = get_db()
        cursor = conn.execute("DELETE FROM adr_reports WHERE id = ?", (report_id,))
        conn.commit()

        if cursor.rowcount == 0:
            return jsonify({"error": "Report not found"}), 404

        return jsonify({"message": f"ADR Report #{report_id} has been deleted."})
    except Exception as e:
        print(f"Error deleting ADR report: {e}")
        return jsonify({"error": "An internal server error occurred."}), 500


# --- NEW: Route to update ADR report status ---
@adr_bp.route("/api/adr-report/<int:report_id>/status", methods=["POST"])
def update_adr_status(report_id):
    if not session.get("admin_id"):
        return jsonify({"error": "Authentication required"}), 401

    data = request.get_json()
    new_status = data.get("status")

    if not new_status:
        return jsonify({"error": "New status is required."}), 400

    try:
        conn = get_db()
        cursor = conn.execute(
            "UPDATE adr_reports SET status = ? WHERE id = ?", (new_status, report_id)
        )
        conn.commit()

        if cursor.rowcount == 0:
            return jsonify({"error": "Report not found"}), 404

        return jsonify(
            {
                "message": f"ADR Report #{report_id} status updated to {new_status}.",
                "new_status": new_status,
            }
        )
    except Exception as e:
        print(f"Error updating ADR status: {e}")
        return jsonify({"error": "An internal server error occurred."}), 500


# --- Helper and Export Routes ---
def get_filtered_adr_data(conn):
    """Helper function to fetch and filter ADR data based on request arguments."""
    adr_search = request.args.get("adr_search", "").strip()
    adr_start = request.args.get("adr_start", "").strip()
    adr_end = request.args.get("adr_end", "").strip()

    query = """
        SELECT
            ar.id, ar.patient_age_range, ar.patient_gender,
            ar.reaction_description, ar.status,
            strftime('%Y-%m-%d', ar.report_date) AS report_date,
            d.name as drug_name, d.batch_number
        FROM adr_reports ar
        JOIN drugs d ON ar.drug_id = d.id
        WHERE 1=1
    """
    params = []

    if adr_search:
        query += " AND (d.name LIKE ? OR d.batch_number LIKE ? OR ar.reaction_description LIKE ?)"
        params.extend([f"%{adr_search}%"] * 3)

    if adr_start and adr_end:
        query += " AND date(ar.report_date) BETWEEN date(?) AND date(?)"
        params.extend([adr_start, adr_end])

    query += " ORDER BY ar.report_date DESC"

    rows = conn.execute(query, params).fetchall()
    return [dict(row) for row in rows]


@adr_bp.route("/admin/adr-reports/export/word")
def export_adr_reports_word():
    if not session.get("admin_id"):
        return jsonify({"error": "Authentication required"}), 401

    conn = get_db()
    reports = get_filtered_adr_data(conn)

    doc = Document()
    doc.add_heading("Adverse Drug Reaction Reports", 0)

    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Drug Name"
    hdr_cells[1].text = "Batch #"
    hdr_cells[2].text = "Patient"
    hdr_cells[3].text = "Reaction"
    hdr_cells[4].text = "Status"
    hdr_cells[5].text = "Report Date"

    for report in reports:
        cells = table.add_row().cells
        cells[0].text = report["drug_name"]
        cells[1].text = report["batch_number"]
        cells[2].text = (
            f"{report['patient_age_range'] or 'N/A'} / {report['patient_gender'] or 'N/A'}"
        )
        cells[3].text = report["reaction_description"]
        cells[4].text = report["status"] or "New"
        cells[5].text = report["report_date"]

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        download_name="adr_reports.docx",
        as_attachment=True,
    )


@adr_bp.route("/admin/adr-reports/export/pdf")
def export_adr_reports_pdf():
    if not session.get("admin_id"):
        return jsonify({"error": "Authentication required"}), 401

    conn = get_db()
    reports = get_filtered_adr_data(conn)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter)
    elements = []
    styles = getSampleStyleSheet()

    elements.append(Paragraph("Adverse Drug Reaction Reports", styles["Title"]))
    elements.append(Spacer(1, 12))

    data = [["Drug", "Batch #", "Patient", "Reaction", "Status", "Report Date"]]
    for report in reports:
        data.append(
            [
                report["drug_name"],
                report["batch_number"],
                f"{report['patient_age_range'] or 'N/A'} / {report['patient_gender'] or 'N/A'}",
                report["reaction_description"],
                report["status"] or "New",
                report["report_date"],
            ]
        )

    table = Table(data, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                ("BACKGROUND", (0, 1), (-1, -1), colors.beige),
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
            ]
        )
    )
    elements.append(table)

    doc.build(elements)
    buf.seek(0)
    return send_file(
        buf, mimetype="application/pdf", download_name="adr_reports.pdf", as_attachment=True
    )