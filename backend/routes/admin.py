from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from datetime import datetime, date, timedelta
from flask import Blueprint, request, send_file, jsonify, url_for, render_template, Response, session, redirect
from sqlite3 import IntegrityError
from backend.models import insert_drug
from backend.database import get_db
import qrcode
import io
import traceback
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from functools import wraps

admin_bp = Blueprint("admin_api", __name__)

def role_required(role):
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if session.get("admin_role") != role:
                return redirect(url_for("admin_login"))
            return f(*args, **kwargs)
        return decorated_function
    return decorator

# =========================
# MODIFIED: Admin Dashboard (Now handles all filtering)
# =========================
@admin_bp.route('/dashboard')
@role_required('regulator')
def admin_dashboard():
    try:
        conn = get_db()

        # --- Counterfeit Reports Filtering Logic ---
        start_date = request.args.get("start", "").strip()
        end_date = request.args.get("end", "").strip()
        filter_type = request.args.get("filter", "").strip()

        counterfeit_query = """
            SELECT r.id, r.drug_name, r.batch_number, r.location, r.note, r.image_filename,
                   strftime('%Y-%m-%d %H:%M:%S', r.reported_on) AS reported_on, r.status,
                   u.email AS submitted_by
            FROM reports r
            LEFT JOIN users u ON r.user_id = u.id
            WHERE 1=1
        """
        counterfeit_params = []

        if filter_type == 'today':
            counterfeit_query += " AND date(r.reported_on) = date('now','localtime')"
        elif start_date and end_date:
            counterfeit_query += " AND date(r.reported_on) BETWEEN date(?) AND date(?)"
            counterfeit_params.extend([start_date, end_date])
        
        counterfeit_query += " ORDER BY r.reported_on DESC, r.id DESC"
        
        counterfeit_rows = conn.execute(counterfeit_query, counterfeit_params).fetchall()
        counterfeit_reports = [dict(r) for r in counterfeit_rows]

        # --- ADR Reports Filtering Logic (Unchanged) ---
        adr_search = request.args.get("adr_search", "").strip()
        adr_start = request.args.get("adr_start", "").strip()
        adr_end = request.args.get("adr_end", "").strip()

        adr_query = """
            SELECT
                ar.id, ar.patient_age_range, ar.patient_gender,
                ar.reaction_description, ar.status,
                strftime('%Y-%m-%d %H:%M:%S', ar.report_date) AS report_date,
                d.name as drug_name, d.batch_number
            FROM adr_reports ar
            JOIN drugs d ON ar.drug_id = d.id
            WHERE 1=1
        """
        adr_params = []

        if adr_search:
            adr_query += " AND (d.name LIKE ? OR d.batch_number LIKE ? OR ar.reaction_description LIKE ?)"
            adr_params.extend([f"%{adr_search}%", f"%{adr_search}%", f"%{adr_search}%"])

        if adr_start and adr_end:
            adr_query += " AND date(ar.report_date) BETWEEN date(?) AND date(?)"
            adr_params.extend([adr_start, adr_end])

        adr_query += " ORDER BY ar.report_date DESC"

        adr_rows = conn.execute(adr_query, adr_params).fetchall()
        adr_reports = [dict(r) for r in adr_rows]

        return render_template(
            'admin.html',
            reports=counterfeit_reports,
            adr_reports=adr_reports,
            current_time=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            adr_search=adr_search,
            adr_start=adr_start,
            adr_end=adr_end
        )
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": f"Server error: {str(e)}"}), 500

# =========================
# Register a new drug batch
# =========================
@admin_bp.post("/register")
@role_required('regulator')
def admin_register():
    if not session.get("admin_id"):
        return jsonify({"error": "Authentication required"}), 401
    try:
        data = request.get_json(silent=True) or request.form.to_dict()
        required = ["name", "batch_number", "mfg_date", "expiry_date", "manufacturer"]
        missing = [f for f in required if not data.get(f)]
        if missing:
            return jsonify({"error": f"Missing fields: {', '.join(missing)}"}), 400
        try:
            insert_drug(
                data["name"],
                data["batch_number"],
                data["mfg_date"],
                data["expiry_date"],
                data["manufacturer"]
            )
        except IntegrityError:
            return jsonify({"error": "Batch number already exists"}), 409
        verify_url = data['batch_number']
        qr_img = qrcode.make(verify_url, box_size=15, border=4)
        buf = io.BytesIO()
        qr_img.save(buf, format="PNG")
        buf.seek(0)
        if request.is_json:
            return send_file(buf, mimetype="image/png", download_name=f"{data['batch_number']}.png")
        import base64
        qr_base64 = base64.b64encode(buf.getvalue()).decode("utf-8")
        conn = get_db()
        rows = conn.execute("""
            SELECT id, drug_name, batch_number, location, note, image_filename,
                   strftime('%Y-%m-%d %H:%M:%S', reported_on) AS reported_on, status
            FROM reports
            ORDER BY reported_on DESC
            LIMIT 20
        """).fetchall()
        reports = [dict(r) for r in rows]
        return render_template("admin.html", qr_image=qr_base64, reports=reports, scroll='qr')
    except Exception as e:
        print("Error in /register:", e)
        traceback.print_exc()
        return jsonify({"error": f"Server error: {str(e)}"}), 500

# =========================
# Registered Drugs Page
# =========================
@admin_bp.get("/drugs")
@role_required('regulator')
def admin_drugs():
    if not session.get("admin_id"):
        return redirect(url_for("admin_login"))
    try:
        conn = get_db()
        search = request.args.get("search", "").strip()
        status = request.args.get("status", "").strip()
        start = request.args.get("start", "").strip()
        end = request.args.get("end", "").strip()
        page = int(request.args.get("page", 1))
        per_page = 20
        offset = (page - 1) * per_page
        base_query = "FROM drugs WHERE 1=1"
        params = []
        if search:
            base_query += " AND (name LIKE ? OR batch_number LIKE ?)"
            params.extend([f"%{search}%", f"%{search}%"])
        if status == "valid":
            base_query += " AND date(expiry_date) >= date('now','localtime')"
        elif status == "expired":
            base_query += " AND date(expiry_date) < date('now','localtime')"
        elif status == "soon":
            base_query += " AND date(expiry_date) BETWEEN date('now','localtime') AND date('now','+30 day','localtime')"
        if start and end:
            base_query += " AND date(created_at) BETWEEN date(?) AND date(?)"
            params.extend([start, end])
        total = conn.execute(f"SELECT COUNT(*) {base_query}", params).fetchone()[0]
        rows = conn.execute(f"""
            SELECT id, name, batch_number, manufacturer, mfg_date, expiry_date, created_at
            {base_query}
            ORDER BY created_at DESC
            LIMIT ? OFFSET ?
        """, params + [per_page, offset]).fetchall()
        total_pages = (total + per_page - 1) // per_page
        return render_template(
            "admin_drugs.html",
            drugs=rows,
            search=search,
            status=status,
            start=start,
            end=end,
            current_date=date.today().isoformat(),
            soon_date=(date.today() + timedelta(days=30)).isoformat(),
            page=page,
            total_pages=total_pages
        )
    except Exception as e:
        print("Error in /drugs:", e)
        traceback.print_exc()
        return jsonify({"error": f"Server error: {str(e)}"}), 500

# =========================
# Delete a registered drug (UPDATED)
# =========================
@admin_bp.delete("/drugs/delete/<int:drug_id>")
@role_required('regulator')
def delete_drug(drug_id):
    if not session.get("admin_id"):
        return jsonify({"error": "Authentication required"}), 401
    
    conn = get_db()
    try:
        conn.execute("DELETE FROM adr_reports WHERE drug_id = ?", (drug_id,))
        cursor = conn.execute("DELETE FROM drugs WHERE id = ?", (drug_id,))
        conn.commit()

        if cursor.rowcount == 0:
            return jsonify({"error": "Drug not found"}), 404
            
        return jsonify({"message": "Drug and all associated reports deleted successfully."})

    except Exception as e:
        conn.rollback() # Roll back the transaction on error
        traceback.print_exc()
        return jsonify({"error": f"Server error: {str(e)}"}), 500


# =========================
# Export Registered Drugs (Word)
# =========================
@admin_bp.get("/drugs/export/word")
@role_required('regulator')
def export_drugs_word():
    if not session.get("admin_id"):
        return redirect(url_for("admin_login"))
    try:
        conn = get_db()
        search = request.args.get("search", "").strip()
        status = request.args.get("status", "").strip()
        start = request.args.get("start", "").strip()
        end = request.args.get("end", "").strip()
        base_query = "FROM drugs WHERE 1=1"
        params = []
        if search:
            base_query += " AND (name LIKE ? OR batch_number LIKE ?)"
            params.extend([f"%{search}%", f"%{search}%"])
        if status == "valid":
            base_query += " AND date(expiry_date) >= date('now','localtime')"
        elif status == "expired":
            base_query += " AND date(expiry_date) < date('now','localtime')"
        elif status == "soon":
            base_query += " AND date(expiry_date) BETWEEN date('now','localtime') AND date('now','+30 day','localtime')"
        if start and end:
            base_query += " AND date(created_at) BETWEEN date(?) AND date(?)"
            params.extend([start, end])
        rows = conn.execute(f"""
            SELECT name, batch_number, manufacturer, mfg_date, expiry_date, created_at
            {base_query}
            ORDER BY created_at DESC
        """, params).fetchall()
        today = date.today()
        soon = today + timedelta(days=30)
        def safe(val):
            return str(val) if val is not None else "N/A"
        doc = Document()
        doc.add_heading("Registered Drugs Report", 0)
        table = doc.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Name"
        hdr_cells[1].text = "Batch Number"
        hdr_cells[2].text = "Manufacturer"
        hdr_cells[3].text = "Mfg Date"
        hdr_cells[4].text = "Expiry Date"
        hdr_cells[5].text = "Status"
        hdr_cells[6].text = "Registered On"
        for row in rows:
            try:
                exp_date = datetime.strptime(
                    safe(row["expiry_date"]), "%Y-%m-%d").date()
            except Exception:
                exp_date = today
            if exp_date < today:
                status_label = "Expired"
            elif exp_date <= soon:
                status_label = "Expiring Soon"
            else:
                status_label = "Valid"
            cells = table.add_row().cells
            cells[0].text = safe(row["name"])
            cells[1].text = safe(row["batch_number"])
            cells[2].text = safe(row["manufacturer"])
            cells[3].text = safe(row["mfg_date"])
            cells[4].text = safe(row["expiry_date"])
            cells[5].text = status_label
            cells[6].text = safe(row["created_at"])
        doc.add_paragraph()
        generated_on = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc.add_paragraph(f"Generated on {generated_on} by Admin System")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(
            buf,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            download_name="registered_drugs.docx",
            as_attachment=True
        )
    except Exception as e:
        print("Error exporting Word:", e)
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

# =========================
# Export Registered Drugs (PDF)
# =========================
@admin_bp.get("/drugs/export/pdf")
@role_required('regulator')
def export_drugs_pdf():
    if not session.get("admin_id"):
        return redirect(url_for("admin_login"))
    try:
        conn = get_db()
        search = request.args.get("search", "").strip()
        status = request.args.get("status", "").strip()
        start = request.args.get("start", "").strip()
        end = request.args.get("end", "").strip()
        base_query = "FROM drugs WHERE 1=1"
        params = []
        if search:
            base_query += " AND (name LIKE ? OR batch_number LIKE ?)"
            params.extend([f"%{search}%", f"%{search}%"])
        if status == "valid":
            base_query += " AND date(expiry_date) >= date('now','localtime')"
        elif status == "expired":
            base_query += " AND date(expiry_date) < date('now','localtime')"
        elif status == "soon":
            base_query += " AND date(expiry_date) BETWEEN date('now','localtime') AND date('now','+30 day','localtime')"
        if start and end:
            base_query += " AND date(created_at) BETWEEN date(?) AND date(?)"
            params.extend([start, end])
        rows = conn.execute(f"""
            SELECT name, batch_number, manufacturer, mfg_date, expiry_date, created_at
            {base_query}
            ORDER BY created_at DESC
        """, params).fetchall()
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=letter)
        elements = []
        styles = getSampleStyleSheet()
        elements.append(Paragraph("Registered Drugs Report", styles["Title"]))
        elements.append(Spacer(1, 12))
        data = [["Name", "Batch Number", "Manufacturer",
                 "Mfg Date", "Expiry Date", "Status", "Registered On"]]
        today = date.today()
        soon = today + timedelta(days=30)
        for row in rows:
            if row["expiry_date"] < today.isoformat():
                status_label = "Expired"
            elif row["expiry_date"] <= soon.isoformat():
                status_label = "Expiring Soon"
            else:
                status_label = "Valid"
            data.append([
                row["name"], row["batch_number"], row["manufacturer"],
                row["mfg_date"], row["expiry_date"], status_label, row["created_at"]
            ])
        table = Table(data, repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f4f4f4")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 11),
            ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("BOX", (0, 0), (-1, -1), 1, colors.black),
        ]))
        elements.append(table)
        elements.append(Spacer(1, 20))
        generated_on = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        footer_text = f"Generated on {generated_on} by Admin System"
        elements.append(Paragraph(footer_text, styles["Normal"]))
        doc.build(elements)
        buf.seek(0)
        return send_file(buf, mimetype="application/pdf",
                         download_name="registered_drugs.pdf", as_attachment=True)
    except Exception as e:
        print("Error exporting PDF:", e)
        traceback.print_exc()
        return jsonify({"error": "Failed to export PDF"}), 500

# =========================
# Reports Count (Unread)
# =========================
@admin_bp.get("/reports/count")
@role_required('regulator')
def reports_count():
    if not session.get("admin_id"):
        return jsonify({"error": "Authentication required"}), 401
    try:
        conn = get_db()
        row = conn.execute("SELECT COUNT(*) as count FROM reports WHERE status = 0").fetchone()
        return jsonify({"count": row["count"] if row else 0})
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": f"Server error: {str(e)}"}), 500

# =========================
# API for Hotspot Map Data
# =========================
@admin_bp.route('/reports/locations')
@role_required('regulator')
def get_report_locations():
    if not session.get("admin_id"):
        return jsonify({"error": "Authentication required"}), 401
    conn = get_db()
    rows = conn.execute("""
        SELECT latitude, longitude, drug_name, batch_number, reported_on
        FROM reports
        WHERE latitude IS NOT NULL AND longitude IS NOT NULL
    """).fetchall()
    locations = [dict(row) for row in rows]
    return jsonify(locations)

# =========================
# Page Rendering Routes
# =========================
@admin_bp.route('/hotspot-map')
@role_required('regulator')
def hotspot_map_page():
    if not session.get("admin_id"):
        return redirect(url_for("admin_login"))
    return render_template('admin_map.html')

@admin_bp.route('/public-db-manager')
@role_required('regulator')
def public_db_manager_page():
    if not session.get("admin_id"):
        return redirect(url_for("admin_login"))
    return render_template('admin_public_db.html')

# =========================
# Session Keep-Alive Endpoint
# =========================
@admin_bp.route('/session/ping', methods=['POST'])
def ping_session():
    """An endpoint for the client to hit to keep the session alive."""
    if not session.get("admin_id"):
        return jsonify({"status": "error", "message": "unauthenticated"}), 401
    # Simply accessing the session is enough to refresh its lifetime
    return jsonify({"status": "ok", "message": "session_refreshed"})